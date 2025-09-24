"""Geracao de relatorios (PDF/DOCX) para avaliacoes RoB 2."""

from io import BytesIO
from typing import Dict, Tuple

from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor

from . import models


def _answer_map() -> Dict[str, str]:
    return {
        "Y": "Sim",
        "PY": "Provavelmente sim",
        "PN": "Provavelmente nao",
        "N": "Nao",
        "NI": "Sem informacao",
        "NA": "Nao se aplica",
    }


# (lower-case prefix) -> (hex fill, hex font)
_JUDGEMENT_COLOR_SCHEME: Dict[str, Tuple[str, str]] = {
    "baixo": ("#22c55e", "#ffffff"),
    "algumas": ("#f97316", "#ffffff"),
    "alto": ("#dc2626", "#ffffff"),
    "critico": ("#7e22ce", "#ffffff"),
}

_DEFAULT_SCHEME = ("#d1d5db", "#111827")
_HEADER_SCHEME = ("#e5e7eb", "#111827")


def _resolve_color(julgamento: str) -> Tuple[colors.Color, colors.Color]:
    normalized = (julgamento or "").strip().lower()
    for prefix, (fill, font) in _JUDGEMENT_COLOR_SCHEME.items():
        if normalized.startswith(prefix):
            return colors.HexColor(fill), colors.HexColor(font)
    return colors.HexColor(_DEFAULT_SCHEME[0]), colors.HexColor(_DEFAULT_SCHEME[1])


def _resolve_hex_colors(julgamento: str) -> Tuple[str, str]:
    normalized = (julgamento or "").strip().lower()
    for prefix, scheme in _JUDGEMENT_COLOR_SCHEME.items():
        if normalized.startswith(prefix):
            return scheme
    return _DEFAULT_SCHEME


def _apply_docx_cell_fill(cell, fill_hex: str, font_hex: str) -> None:
    """Apply background and font colors to a python-docx table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex.lstrip("#").upper())
    tc_pr.append(shd)

    rgb = RGBColor.from_string(font_hex.lstrip("#"))
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = rgb


def _docx_header_style(cell) -> None:
    fill_hex, font_hex = _HEADER_SCHEME
    _apply_docx_cell_fill(cell, fill_hex, font_hex)
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True


def generate_pdf_report(avaliacao: models.Evaluation) -> bytes:
    """Gera um relatorio PDF para uma avaliacao."""
    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=2 * cm,
        leftMargin=2 * cm,
        topMargin=2 * cm,
        bottomMargin=2 * cm,
    )
    styles = getSampleStyleSheet()
    story = []

    estudo = avaliacao.resultado.estudo
    resultado = avaliacao.resultado

    story.append(Paragraph("Relatorio de Avaliacao RoB 2", styles["Title"]))
    story.append(Spacer(1, 0.3 * cm))
    story.append(Paragraph(f"Artigo/Referencia: {estudo.referencia}", styles["Normal"]))
    story.append(Paragraph(f"Dominio do estudo: {estudo.desenho or '-'}", styles["Normal"]))
    story.append(Paragraph(f"Desfecho avaliado: {resultado.desfecho}", styles["Normal"]))
    story.append(Spacer(1, 0.5 * cm))

    # Dashboard resumido
    story.append(Paragraph("Resumo visual do risco de vies", styles["Heading2"]))
    dashboard_rows = [["Escopo", "Julgamento"]]
    dashboard_styles = [
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(_HEADER_SCHEME[0])),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor(_HEADER_SCHEME[1])),
        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
    ]

    global_judgement = avaliacao.julgamento_global or "Nao calculado"
    dashboard_rows.append(["Global", global_judgement])

    dominios = sorted(avaliacao.dominios, key=lambda dom: dom.tipo)
    for dom in dominios:
        dashboard_rows.append([f"Dominio {dom.tipo}", dom.julgamento or "Nao avaliado"])

    for row_idx in range(1, len(dashboard_rows)):
        judgement = dashboard_rows[row_idx][1]
        fill_color, font_color = _resolve_color(judgement)
        dashboard_styles.append(("BACKGROUND", (1, row_idx), (1, row_idx), fill_color))
        dashboard_styles.append(("TEXTCOLOR", (1, row_idx), (1, row_idx), font_color))
        dashboard_styles.append(("FONTNAME", (1, row_idx), (1, row_idx), "Helvetica-Bold"))

    dashboard_table = Table(dashboard_rows, colWidths=[6 * cm, 6 * cm])
    dashboard_table.setStyle(TableStyle(dashboard_styles))
    story.append(dashboard_table)
    story.append(Spacer(1, 0.2 * cm))
    story.append(Paragraph("Legenda: Baixo (verde), Algumas preocupacoes (laranja), Alto (vermelho).", styles["Italic"]))
    story.append(Spacer(1, 0.4 * cm))

    if avaliacao.pre_consideracoes:
        story.append(Paragraph("Pre-consideracoes", styles["Heading2"]))
        story.append(Paragraph(avaliacao.pre_consideracoes, styles["Normal"]))
        story.append(Spacer(1, 0.4 * cm))

    ans_map = _answer_map()
    for dom in dominios:
        story.append(Paragraph(f"Dominio {dom.tipo}", styles["Heading2"]))

        itens = [["Item", "Resposta", "Observacao"]]
        obs_map = dom.observacoes_itens or {}
        for pergunta_id, resposta in (dom.respostas or {}).items():
            itens.append(
                [
                    pergunta_id,
                    ans_map.get(resposta, resposta),
                    obs_map.get(pergunta_id) or "-",
                ]
            )
        table = Table(itens, colWidths=[3 * cm, 4 * cm, 9 * cm])
        table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(_HEADER_SCHEME[0])),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor(_HEADER_SCHEME[1])),
                    ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ]
            )
        )
        story.append(table)
        story.append(Spacer(1, 0.2 * cm))
        story.append(Paragraph(f"Julgamento do dominio: {dom.julgamento or '-'}", styles["Italic"]))
        if dom.justificativa:
            story.append(Paragraph(dom.justificativa, styles["Normal"]))
        story.append(Spacer(1, 0.4 * cm))

    story.append(Paragraph("Julgamento global", styles["Heading2"]))
    story.append(Paragraph(global_judgement, styles["Normal"]))
    if avaliacao.justificativa_global:
        story.append(Paragraph(avaliacao.justificativa_global, styles["Normal"]))
    if avaliacao.direcao_global:
        story.append(Paragraph(f"Direcao do vies: {avaliacao.direcao_global.value}", styles["Normal"]))

    doc.build(story)
    pdf_bytes = buffer.getvalue()
    buffer.close()
    return pdf_bytes


def generate_docx_report(avaliacao: models.Evaluation) -> bytes:
    """Generate a DOCX narrative report for an evaluation."""
    buffer = BytesIO()
    document = Document()

    estudo = avaliacao.resultado.estudo
    resultado = avaliacao.resultado

    document.add_heading("Relatorio de Avaliacao RoB 2", level=1)
    document.add_paragraph(f"Artigo/Referencia: {estudo.referencia}")
    document.add_paragraph(f"Dominio do estudo: {estudo.desenho or '-'}")
    document.add_paragraph(f"Desfecho avaliado: {resultado.desfecho}")

    # Dashboard
    document.add_heading("Resumo visual do risco de vies", level=2)
    dashboard_table = document.add_table(rows=1, cols=2)
    header_cells = dashboard_table.rows[0].cells
    header_cells[0].text = "Escopo"
    header_cells[1].text = "Julgamento"
    for cell in header_cells:
        _docx_header_style(cell)

    global_judgement = avaliacao.julgamento_global or "Nao calculado"
    row = dashboard_table.add_row().cells
    row[0].text = "Global"
    row[1].text = global_judgement
    fill_hex, font_hex = _resolve_hex_colors(global_judgement)
    _apply_docx_cell_fill(row[1], fill_hex, font_hex)

    for dom in sorted(avaliacao.dominios, key=lambda item: item.tipo):
        row = dashboard_table.add_row().cells
        row[0].text = f"Dominio {dom.tipo}"
        row[1].text = dom.julgamento or "Nao avaliado"
        fill_hex, font_hex = _resolve_hex_colors(dom.julgamento or "")
        _apply_docx_cell_fill(row[1], fill_hex, font_hex)

    legend_paragraph = document.add_paragraph()
    legend_run = legend_paragraph.add_run("Legenda: Baixo (verde), Algumas preocupacoes (laranja), Alto (vermelho).")
    legend_run.italic = True

    if avaliacao.pre_consideracoes:
        document.add_heading("Pre-consideracoes", level=2)
        document.add_paragraph(avaliacao.pre_consideracoes)

    ans_map = _answer_map()
    for dom in sorted(avaliacao.dominios, key=lambda item: item.tipo):
        document.add_heading(f"Dominio {dom.tipo}", level=2)
        table = document.add_table(rows=1, cols=3)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Item"
        hdr_cells[1].text = "Resposta"
        hdr_cells[2].text = "Observacao"
        for cell in hdr_cells:
            _docx_header_style(cell)
        obs_map = dom.observacoes_itens or {}
        for pergunta_id, resposta in (dom.respostas or {}).items():
            row_cells = table.add_row().cells
            row_cells[0].text = str(pergunta_id)
            row_cells[1].text = ans_map.get(resposta, resposta or "-")
            row_cells[2].text = str(obs_map.get(pergunta_id) or "-")
        document.add_paragraph(f"Julgamento do dominio: {dom.julgamento or '-'}")
        if dom.justificativa:
            document.add_paragraph(dom.justificativa)
        document.add_paragraph("")

    document.add_heading("Julgamento global", level=2)
    document.add_paragraph(global_judgement)
    if avaliacao.justificativa_global:
        document.add_paragraph(avaliacao.justificativa_global)
    if avaliacao.direcao_global:
        document.add_paragraph(f"Direcao do vies: {avaliacao.direcao_global.value}")

    document.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
